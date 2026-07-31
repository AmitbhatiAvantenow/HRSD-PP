"""Renders a template's canvas_data (see document_template.py for the JSON shape)
into PDF bytes (reportlab, true absolute x/y positioning) or DOCX bytes
(python-docx, flow-based -- blocks are laid out top-to-bottom since Word has no
native absolute-positioning surface without complex floating-shape XML; this is
a documented limitation, not a bug).

All coordinates are in points (pt) -- reportlab and python-docx both accept pt
natively, so no unit conversion happens between the two renderers.
"""
import base64
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from jinja2.sandbox import SandboxedEnvironment
from reportlab.lib import colors as rl_colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import Paragraph, Table, TableStyle

from odoo import fields
from odoo.exceptions import ValidationError
from odoo.tools.translate import _

PAGE_SIZES_PT = {
    'a4': (595, 842),
    'letter': (612, 792),
    'legal': (612, 1008),
}


def page_dims(paper_size, orientation):
    w, h = PAGE_SIZES_PT.get(paper_size, PAGE_SIZES_PT['a4'])
    return (h, w) if orientation == 'landscape' else (w, h)


# ---------------------------------------------------------------------------
# Variable substitution (sandboxed -- block text is user-authored content
# rendered against runtime data, so a plain Environment would let a crafted
# {{ }} expression reach __class__/__globals__ style escapes)
# ---------------------------------------------------------------------------

_JINJA_ENV = SandboxedEnvironment(autoescape=False)


def render_text(text, variables):
    if not text or '{{' not in text:
        return text or ''
    try:
        return _JINJA_ENV.from_string(text).render(**variables)
    except Exception:
        # Fail open: a malformed placeholder shouldn't abort the whole document.
        return text


def _substitute_value(value, variables):
    if isinstance(value, str):
        return render_text(value, variables)
    if isinstance(value, list):
        return [_substitute_value(v, variables) for v in value]
    if isinstance(value, dict):
        return {k: _substitute_value(v, variables) for k, v in value.items()}
    return value


def substitute_block_props(block, variables):
    new_block = dict(block)
    new_block['props'] = _substitute_value(block.get('props') or {}, variables)
    return new_block


def coerce_variable_values(template, raw_values):
    """Turn {key: raw user input} into {key: display-ready primitive} using each
    variable's declared type, raising for missing required values."""
    raw_values = raw_values or {}
    variables = {}
    for v in template.variable_ids:
        val = raw_values.get(v.key, v.default_value or '')
        if v.is_required and val in (None, ''):
            raise ValidationError(_('Variable "%s" is required.') % v.name)
        if v.variable_type == 'currency':
            try:
                val = f'{float(val or 0):,.2f}'
            except (TypeError, ValueError):
                val = str(val)
        elif v.variable_type == 'number':
            val = str(val)
        elif v.variable_type == 'date' and val:
            try:
                d = fields.Date.from_string(val) if isinstance(val, str) else val
                val = d.strftime('%d %b %Y')
            except Exception:
                val = str(val)
        elif v.variable_type == 'boolean':
            val = 'Yes' if val in (True, 'true', 'True', '1', 1) else 'No'
        variables[v.key] = val
    return variables


# ---------------------------------------------------------------------------
# Page-break partitioning: page_break blocks are horizontal markers (only `y`
# matters) that split the one continuous canvas into page-groups, each
# re-based so its own top is y=0.
# ---------------------------------------------------------------------------

def partition_pages(blocks, page_h_pt):
    breaks = sorted(b['y'] for b in blocks if b.get('type') == 'page_break')
    boundaries = [0] + breaks + [float('inf')]
    pages = [[] for _ in range(len(boundaries) - 1)]
    for b in blocks:
        if b.get('type') == 'page_break':
            continue
        y = b.get('y', 0)
        for i in range(len(boundaries) - 1):
            if boundaries[i] <= y < boundaries[i + 1]:
                nb = dict(b)
                nb['y'] = y - boundaries[i]
                pages[i].append(nb)
                break
    return pages or [[]]


def _decode_image_data(data_str):
    if not data_str:
        return None
    if isinstance(data_str, str) and data_str.strip().startswith('data:') and ',' in data_str:
        data_str = data_str.split(',', 1)[1]
    try:
        return base64.b64decode(data_str)
    except Exception:
        return None


def _make_qr_png(data):
    import qrcode
    img = qrcode.make(data or '')
    buf = BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def _make_barcode_png(data, fmt='code128'):
    import barcode
    from barcode.writer import ImageWriter
    try:
        barcode_cls = barcode.get_barcode_class(fmt or 'code128')
        obj = barcode_cls(data or '000000000', writer=ImageWriter())
    except Exception:
        obj = barcode.get_barcode_class('code128')(data or '000000000', writer=ImageWriter())
    buf = BytesIO()
    obj.write(buf, options={'write_text': False})
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF renderer (reportlab) -- canvas origin is bottom-left; canvas_data's
# origin is top-left (natural for an HTML/CSS-based builder), so every block
# converts once via _to_pdf_y().
# ---------------------------------------------------------------------------

def _to_pdf_y(page_h, y, h):
    return page_h - y - h


ALIGN_MAP_PDF = {'left': TA_LEFT, 'center': TA_CENTER, 'right': TA_RIGHT}


def _font_name(bold, italic):
    if bold and italic:
        return 'Helvetica-BoldOblique'
    if bold:
        return 'Helvetica-Bold'
    if italic:
        return 'Helvetica-Oblique'
    return 'Helvetica'


def _text_style(props, block_type):
    level = props.get('level', 1)
    default_size = {1: 24, 2: 18, 3: 14}.get(level, 14) if block_type == 'heading' else 11
    font_size = props.get('font_size') or default_size
    bold = props.get('bold', block_type == 'heading')
    italic = props.get('italic', False)
    color = props.get('color') or '#111111'
    align = ALIGN_MAP_PDF.get(props.get('align', 'left'), TA_LEFT)
    return ParagraphStyle(
        name='block', fontName=_font_name(bold, italic), fontSize=font_size,
        leading=font_size * 1.25, textColor=HexColor(color), alignment=align,
    )


def measure_text_height(text, props, block_type, width_pt):
    """Real wrapped height (pt) for a text/heading block at a given width, using the
    exact same style/Paragraph logic _render_text_block_pdf draws with -- so a layout
    computed from this (e.g. stacking imported docx paragraphs) never overlaps what
    actually gets rendered, regardless of how many lines the text wraps to."""
    text = (text or '').replace('\n', '<br/>')
    if not text:
        return 0
    style = _text_style(props, block_type)
    _, h = Paragraph(text, style).wrap(width_pt, 10000)
    return h


def _render_text_block_pdf(c, block, page_h):
    props = block.get('props') or {}
    text = (props.get('text') or '').replace('\n', '<br/>')
    if not text:
        return
    style = _text_style(props, block.get('type'))
    x, y, w = block['x'], block['y'], block['w']
    p = Paragraph(text, style)
    _, h_used = p.wrap(w, 10000)
    p.drawOn(c, x, page_h - y - h_used)


def _draw_image_pdf(c, image_bytes, x, y, w, h, page_h, fit='contain'):
    if not image_bytes:
        return
    try:
        img = ImageReader(BytesIO(image_bytes))
    except Exception:
        return
    c.drawImage(img, x, _to_pdf_y(page_h, y, h), width=w, height=h,
                preserveAspectRatio=(fit == 'contain'), mask='auto')


def _render_image_block_pdf(c, block, page_h):
    props = block.get('props') or {}
    image_bytes = _decode_image_data(props.get('image_data'))
    _draw_image_pdf(c, image_bytes, block['x'], block['y'], block['w'], block['h'], page_h, props.get('fit', 'contain'))


def _render_qr_block_pdf(c, block, page_h):
    props = block.get('props') or {}
    png = _make_qr_png(props.get('data', ''))
    _draw_image_pdf(c, png, block['x'], block['y'], block['w'], block['h'], page_h)


def _render_barcode_block_pdf(c, block, page_h):
    props = block.get('props') or {}
    try:
        png = _make_barcode_png(props.get('data', ''), props.get('format', 'code128'))
    except Exception:
        png = None
    if png:
        _draw_image_pdf(c, png, block['x'], block['y'], block['w'], block['h'], page_h)


def _render_table_pdf(c, block, page_h):
    props = block.get('props') or {}
    headers = props.get('headers') or []
    rows = props.get('rows') or []
    data = ([headers] if headers else []) + rows
    if not data:
        return
    x, y, w, h = block['x'], block['y'], block['w'], block['h']
    n_cols = max(len(row) for row in data)
    t = Table(data, colWidths=[w / n_cols] * n_cols)
    style_cmds = [
        ('GRID', (0, 0), (-1, -1), 0.75, rl_colors.HexColor('#d8dee6')),
        ('FONTSIZE', (0, 0), (-1, -1), props.get('font_size', 9)),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]
    if headers:
        style_cmds += [
            ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#eef2f6')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ]
    t.setStyle(TableStyle(style_cmds))
    _, h_used = t.wrap(w, h or 10000)
    t.drawOn(c, x, page_h - y - h_used)


def _render_divider_pdf(c, block, page_h):
    props = block.get('props') or {}
    x, y, w = block['x'], block['y'], block['w']
    c.setStrokeColor(HexColor(props.get('color') or '#cccccc'))
    c.setLineWidth(props.get('thickness', 1))
    c.setDash([4, 3] if props.get('style') == 'dashed' else [])
    yy = page_h - y
    c.line(x, yy, x + w, yy)
    c.setDash([])


def _render_signature_pdf(c, block, page_h):
    props = block.get('props') or {}
    x, y, w, h = block['x'], block['y'], block['w'], block['h']
    c.setStrokeColor(HexColor(props.get('border_color') or '#999999'))
    c.setDash([3, 2] if props.get('box_style') == 'dashed' else [])
    c.rect(x, page_h - y - h, w, h)
    c.setDash([])
    image_bytes = _decode_image_data(props.get('image_data'))
    if image_bytes:
        _draw_image_pdf(c, image_bytes, x + 4, y + 4, w - 8, max(h * 0.6, 10), page_h)
    c.setFont('Helvetica', 9)
    c.setFillColor(HexColor('#555555'))
    c.drawString(x + 4, page_h - y - h + 4, props.get('label') or 'Signature')


def _render_shape_pdf(c, block, page_h):
    props = block.get('props') or {}
    x, y, w, h = block['x'], block['y'], block['w'], block['h']
    shape = props.get('shape', 'rectangle')
    fill = props.get('fill_color')
    c.setStrokeColor(HexColor(props.get('border_color') or '#333333'))
    if fill:
        c.setFillColor(HexColor(fill))
    if shape == 'circle':
        r = min(w, h) / 2
        c.circle(x + w / 2, page_h - y - h / 2, r, stroke=1, fill=1 if fill else 0)
    elif shape == 'line':
        yy = page_h - y - h / 2
        c.line(x, yy, x + w, yy)
    else:
        c.rect(x, page_h - y - h, w, h, stroke=1, fill=1 if fill else 0)


def _render_icon_pdf(c, block, page_h):
    props = block.get('props') or {}
    x, y, w, h = block['x'], block['y'], block['w'], block['h']
    r = min(w, h) / 2
    c.setFillColor(HexColor(props.get('color') or '#2f8f5b'))
    c.circle(x + r, page_h - y - h + r, r, stroke=0, fill=1)


def _render_chart_pdf(c, block, page_h):
    props = block.get('props') or {}
    x, y, w, h = block['x'], block['y'], block['w'], block['h']
    data = props.get('data') or []
    if not data:
        return
    max_v = max((d.get('value', 0) for d in data), default=1) or 1
    n = len(data)
    gap = 6
    bar_w = max(4, (w - gap * (n + 1)) / n)
    base_y = page_h - y - h
    for i, d in enumerate(data):
        bar_h = max(2, (d.get('value', 0) / max_v) * (h - 14))
        bx = x + gap + i * (bar_w + gap)
        c.setFillColor(HexColor(props.get('color') or '#2f8f5b'))
        c.rect(bx, base_y + 14, bar_w, bar_h, stroke=0, fill=1)
        c.setFillColor(HexColor('#555555'))
        c.setFont('Helvetica', 7)
        c.drawCentredString(bx + bar_w / 2, base_y + 3, str(d.get('label', ''))[:8])


_PDF_RENDERERS = {
    'text': _render_text_block_pdf,
    'heading': _render_text_block_pdf,
    'image': _render_image_block_pdf,
    'logo': _render_image_block_pdf,
    'stamp': _render_image_block_pdf,
    'table': _render_table_pdf,
    'qr': _render_qr_block_pdf,
    'barcode': _render_barcode_block_pdf,
    'divider': _render_divider_pdf,
    'signature': _render_signature_pdf,
    'shape': _render_shape_pdf,
    'icon': _render_icon_pdf,
    'chart': _render_chart_pdf,
}


def render_pdf(page_meta, pages):
    buf = BytesIO()
    page_w = page_meta.get('width_pt', 595)
    page_h = page_meta.get('height_pt', 842)
    margin = page_meta.get('margin_pt', 36)
    c = pdfcanvas.Canvas(buf, pagesize=(page_w, page_h))
    for page_blocks in pages:
        bg = page_meta.get('background_color')
        if bg:
            c.setFillColor(HexColor(bg))
            c.rect(0, 0, page_w, page_h, stroke=0, fill=1)
        header = page_meta.get('header_text')
        if header:
            c.setFont('Helvetica', 8)
            c.setFillColor(HexColor('#888888'))
            c.drawString(margin, page_h - margin / 2, header)
        footer = page_meta.get('footer_text')
        if footer:
            c.setFont('Helvetica', 8)
            c.setFillColor(HexColor('#888888'))
            c.drawString(margin, margin / 2, footer)
        for block in sorted(page_blocks, key=lambda b: b.get('z', 0)):
            renderer = _PDF_RENDERERS.get(block.get('type'))
            if not renderer:
                continue
            try:
                renderer(c, block, page_h)
            except Exception:
                continue
        c.showPage()
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX renderer (python-docx) -- flow-based: blocks are laid out top-to-bottom
# sorted by (y, x), NOT at their absolute canvas position. Documented
# limitation, not a bug: Word has no simple absolute-positioning surface.
# ---------------------------------------------------------------------------

ALIGN_MAP_DOCX = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
}


def _set_cell_background(cell, hex_color):
    if not hex_color:
        return
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shd)


def _set_paragraph_bottom_border(paragraph, color='cccccc', size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_border(cell, color='999999'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def _add_text_block_docx(doc, block):
    props = block.get('props') or {}
    text = props.get('text') or ''
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = ALIGN_MAP_DOCX.get(props.get('align', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    is_heading = block.get('type') == 'heading'
    run.font.size = Pt(props.get('font_size') or (20 if is_heading else 11))
    run.bold = props.get('bold', is_heading)
    run.italic = props.get('italic', False)
    try:
        run.font.color.rgb = RGBColor.from_string((props.get('color') or '#111111').lstrip('#'))
    except Exception:
        pass


def _add_image_bytes_docx(doc, image_bytes, width_pt):
    if not image_bytes:
        return
    try:
        doc.add_picture(BytesIO(image_bytes), width=Pt(max(width_pt or 120, 20)))
    except Exception:
        pass


def _add_image_block_docx(doc, block):
    props = block.get('props') or {}
    _add_image_bytes_docx(doc, _decode_image_data(props.get('image_data')), block.get('w'))


def _add_qr_block_docx(doc, block):
    props = block.get('props') or {}
    _add_image_bytes_docx(doc, _make_qr_png(props.get('data', '')), min(block.get('w', 100), 150))


def _add_barcode_block_docx(doc, block):
    props = block.get('props') or {}
    try:
        png = _make_barcode_png(props.get('data', ''), props.get('format', 'code128'))
    except Exception:
        png = None
    _add_image_bytes_docx(doc, png, min(block.get('w', 150), 220))


def _add_table_block_docx(doc, block):
    props = block.get('props') or {}
    headers = props.get('headers') or []
    rows = props.get('rows') or []
    n_cols = len(headers) or (len(rows[0]) if rows else 0)
    if not n_cols:
        return
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'
    if headers:
        row_cells = table.add_row().cells
        for i, val in enumerate(headers):
            row_cells[i].text = str(val)
            for run in row_cells[i].paragraphs[0].runs:
                run.bold = True
            _set_cell_background(row_cells[i], 'eef2f6')
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = str(val)


def _add_divider_block_docx(doc, block):
    props = block.get('props') or {}
    _set_paragraph_bottom_border(doc.add_paragraph(), color=(props.get('color') or '#cccccc').lstrip('#'))


def _add_signature_block_docx(doc, block):
    props = block.get('props') or {}
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_border(cell)
    image_bytes = _decode_image_data(props.get('image_data'))
    p = cell.paragraphs[0]
    if image_bytes:
        try:
            run = p.add_run()
            run.add_picture(BytesIO(image_bytes), width=Pt(min(block.get('w', 150), 150)))
        except Exception:
            pass
    cell.add_paragraph(props.get('label') or 'Signature')


def _add_shape_or_icon_block_docx(doc, block):
    # Documented limitation: no true shapes/icons in flow-based docx output --
    # rendered as a small bordered placeholder cell instead.
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_border(cell)
    cell.text = (block.get('type') or '').title()


def _add_chart_block_docx(doc, block):
    props = block.get('props') or {}
    data = props.get('data') or []
    if not data:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = 'Label', 'Value'
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for d in data:
        cells = table.add_row().cells
        cells[0].text = str(d.get('label', ''))
        cells[1].text = str(d.get('value', ''))


_DOCX_RENDERERS = {
    'text': _add_text_block_docx,
    'heading': _add_text_block_docx,
    'image': _add_image_block_docx,
    'logo': _add_image_block_docx,
    'stamp': _add_image_block_docx,
    'table': _add_table_block_docx,
    'qr': _add_qr_block_docx,
    'barcode': _add_barcode_block_docx,
    'divider': _add_divider_block_docx,
    'signature': _add_signature_block_docx,
    'shape': _add_shape_or_icon_block_docx,
    'icon': _add_shape_or_icon_block_docx,
    'chart': _add_chart_block_docx,
}


def render_docx(page_meta, pages):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Pt(page_meta.get('width_pt', 595))
    section.page_height = Pt(page_meta.get('height_pt', 842))

    header_text = page_meta.get('header_text')
    if header_text:
        section.header.paragraphs[0].text = header_text
    footer_text = page_meta.get('footer_text')
    if footer_text:
        section.footer.paragraphs[0].text = footer_text

    for page_index, page_blocks in enumerate(pages):
        if page_index > 0:
            doc.add_page_break()
        ordered = sorted(page_blocks, key=lambda b: (round(b.get('y', 0)), round(b.get('x', 0)), b.get('z', 0)))
        for block in ordered:
            renderer = _DOCX_RENDERERS.get(block.get('type'))
            if not renderer:
                continue
            try:
                renderer(doc, block)
            except Exception:
                continue

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
